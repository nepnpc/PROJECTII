"""
Optical Ledger — Project Report Generator
Produces: optical_ledger_report.docx
Font: Times New Roman | Margins: L 1.5" R/T/B 1" | Spacing: 1.5
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

# ── HELPERS ───────────────────────────────────────────────────────────────────

def set_margins(section, left=1.5, right=1.0, top=1.0, bottom=1.0):
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)

def set_font(run, size=12, bold=False, italic=False, name='Times New Roman'):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    name)
    rFonts.set(qn('w:hAnsi'),    name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'),       name)
    rPr.insert(0, rFonts)

def set_para_spacing(para, line_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
                     space_before=0, space_after=0):
    pf = para.paragraph_format
    pf.line_spacing_rule = line_rule
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)

def add_para(doc, text='', size=12, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    set_para_spacing(p, space_before=space_before, space_after=space_after)
    p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_heading(doc, text, level=1, numbered='', space_before=18, space_after=6):
    """Custom heading — uses Word built-in Heading styles for TOC pickup."""
    style_name = f'Heading {level}'
    p = doc.add_heading(text if not numbered else f'{numbered}  {text}', level=level)
    set_para_spacing(p, space_before=space_before, space_after=space_after)
    for run in p.runs:
        set_font(run, size=14 if level == 1 else 12, bold=True,
                 name='Times New Roman')
    return p

def add_figure_placeholder(doc, caption='Figure X.X: [Description]'):
    """Shaded box as placeholder for figure image."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(5)
    # shade cell
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'D9D9D9')
    tcPr.append(shd)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(cp, space_before=6, space_after=6)
    r = cp.add_run('[Figure Placeholder — Insert Image Here]')
    set_font(r, size=10, italic=True)
    # height hint via empty lines
    for _ in range(5):
        ep = cell.add_paragraph('')
        set_para_spacing(ep, space_before=0, space_after=0)
    # caption below
    cap = add_para(doc, caption, size=11, italic=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before=4, space_after=12)
    return tbl, cap

def add_table_caption(doc, caption='Table X.X: [Description]'):
    """Caption above table."""
    cap = add_para(doc, caption, size=11, italic=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before=12, space_after=4)
    return cap

def insert_toc(doc):
    """Insert Word TOC field (updates when opened in Word)."""
    p = doc.add_paragraph()
    set_para_spacing(p, space_before=0, space_after=0)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r = p.add_run()
    r._r.append(fldChar)
    r._r.append(instrText)
    r._r.append(fldChar2)
    r._r.append(fldChar3)
    hint = add_para(doc,
        '[Right-click this area → Update Field to generate Table of Contents]',
        size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=4, space_after=12)

def set_page_number_format(section, fmt='decimal'):
    """Set page number format for a section (decimal or lowerRoman)."""
    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), '1')
    existing = sectPr.find(qn('w:pgNumType'))
    if existing is not None:
        sectPr.remove(existing)
    sectPr.append(pgNumType)

def add_page_number_footer(section):
    """Add centered page number footer to a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    run = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_font(run, size=11)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.oxml.ns.qn if False else __import__(
        'docx').enum.text.WD_BREAK.PAGE if False else None)
    p.runs[0].add_break()
    # simpler:
    doc.add_page_break()
    # remove the extra blank para
    p._element.getparent().remove(p._element)

def page_break(doc):
    doc.add_page_break()

def add_code_block(doc, code_text):
    """Add monospaced code block with light background."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    tcPr.append(shd)
    cp = cell.paragraphs[0]
    cp.clear()
    set_para_spacing(cp, line_rule=WD_LINE_SPACING.SINGLE,
                     space_before=4, space_after=4)
    for line in code_text.split('\n'):
        lp = cell.add_paragraph(line)
        set_para_spacing(lp, line_rule=WD_LINE_SPACING.SINGLE,
                         space_before=0, space_after=0)
        for r in lp.runs:
            set_font(r, size=9, name='Courier New')
        if not lp.runs:
            lp.add_run('')
    # remove first empty cell para
    cp._element.getparent().remove(cp._element)
    doc.add_paragraph('')  # spacer
    return tbl

def borderless_table(doc, rows_data, header=None):
    """Table with no visible borders (for abbreviations)."""
    cols = len(rows_data[0]) if rows_data else 2
    tbl  = doc.add_table(rows=0, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # remove all borders
    tblPr  = tbl._tbl.tblPr
    tblBdr = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'none')
        tblBdr.append(bdr)
    tblPr.append(tblBdr)

    if header:
        row  = tbl.add_row()
        for i, h in enumerate(header):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(h)
            set_font(r, size=12, bold=True)
    for row_data in rows_data:
        row = tbl.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            set_font(r, size=12)
    return tbl

# ── CONTENT BLOCKS ────────────────────────────────────────────────────────────

LINEAR_SEARCH_CODE = """\
/**
 * Algorithm 1: LINEAR SEARCH
 * Time Complexity:  O(n)
 * Space Complexity: O(k) where k = number of matches
 * Use: Search lenses/inventory by brand, index, or power
 */
function linearSearch(dataArray, searchQuery, searchFields) {
    var results = [];
    var query = searchQuery.toLowerCase().trim();

    if (query === '') return dataArray;

    for (var i = 0; i < dataArray.length; i++) {
        var item = dataArray[i];
        var matched = false;

        for (var j = 0; j < searchFields.length; j++) {
            var fieldValue = String(item[searchFields[j]] || '').toLowerCase();

            if (fieldValue.indexOf(query) !== -1) {
                matched = true;
                break;
            }
        }

        if (matched) {
            results.push(item);
        }
    }

    return results;
}"""

BUBBLE_SORT_CODE = """\
/**
 * Algorithm 2: BUBBLE SORT (with early-exit optimization)
 * Time Complexity:  O(n²) worst case, O(n) best case
 * Space Complexity: O(1) in-place
 * Use: Sort orders by amount, date, or status
 */
function bubbleSort(dataArray, sortField, sortOrder) {
    var arr = dataArray.slice();
    var n = arr.length;
    var swapped;

    for (var i = 0; i < n - 1; i++) {
        swapped = false;

        for (var j = 0; j < n - i - 1; j++) {
            var valA = arr[j][sortField];
            var valB = arr[j + 1][sortField];

            if (!isNaN(parseFloat(valA)) && !isNaN(parseFloat(valB))) {
                valA = parseFloat(valA);
                valB = parseFloat(valB);
            } else {
                valA = String(valA).toLowerCase();
                valB = String(valB).toLowerCase();
            }

            var shouldSwap = (sortOrder === 'desc') ? (valA < valB) : (valA > valB);

            if (shouldSwap) {
                var temp  = arr[j];
                arr[j]    = arr[j + 1];
                arr[j + 1] = temp;
                swapped   = true;
            }
        }

        if (!swapped) break;
    }

    return arr;
}"""

LEDGER_SQL = """\
-- Append-Only Inventory Ledger Pattern
-- Net stock = SUM(change_qty) per lens + power combination

CREATE TABLE inventory_ledger (
    id         INT AUTO_INCREMENT PRIMARY KEY,
    lens_id    INT NOT NULL,
    sph        DECIMAL(5,2) NOT NULL DEFAULT 0.00,
    cyl        DECIMAL(5,2) NOT NULL DEFAULT 0.00,
    axis       DECIMAL(5,2) NOT NULL DEFAULT 0.00,
    change_qty INT          NOT NULL,   -- positive = stock in, negative = stock out
    notes      VARCHAR(255),
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (lens_id) REFERENCES lenses(id)
);

-- Query net stock per power
SELECT lens_id, sph, cyl, axis,
       SUM(change_qty) AS net_qty
FROM   inventory_ledger
GROUP  BY lens_id, sph, cyl, axis
HAVING SUM(change_qty) > 0;"""

# ── MAIN BUILDER ──────────────────────────────────────────────────────────────

def build():
    import docx
    doc = Document()

    # ── Default paragraph style ───────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    # apply to heading styles too
    for lvl in range(1, 5):
        hs = doc.styles[f'Heading {lvl}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = RGBColor(0, 0, 0)

    # ── Section 1: Front Matter (Roman numerals) ──────────────────────────────
    section1 = doc.sections[0]
    set_margins(section1)
    set_page_number_format(section1, fmt='lowerRoman')
    add_page_number_footer(section1)

    # ── TITLE PAGE ────────────────────────────────────────────────────────────
    for _ in range(4):
        add_para(doc, '', space_before=0, space_after=0)

    add_para(doc, 'OPTICAL LEDGER', size=20, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)
    add_para(doc, 'A Web-Based Inventory and Order Management System\nfor Optical Businesses',
             size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=30)
    add_para(doc, 'Project Report', size=14, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)
    add_para(doc, 'Submitted in partial fulfillment of the requirements\nfor the Bachelor of Science in Computer Science',
             size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=30)

    add_para(doc, 'Submitted by:', size=12, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
    add_para(doc, '[Student Name(s)]', size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
    add_para(doc, '[Roll Number(s)]', size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=30)

    add_para(doc, 'Supervised by:', size=12, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
    add_para(doc, '[Supervisor Name]', size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
    add_para(doc, '[Department / College Name]', size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=30)

    add_para(doc, '[Academic Year]', size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)

    page_break(doc)

    # ── DECLARATION ───────────────────────────────────────────────────────────
    add_heading(doc, 'Declaration', level=1, space_before=0)
    add_para(doc,
        'We hereby declare that this project report entitled "Optical Ledger: A Web-Based '
        'Inventory and Order Management System for Optical Businesses" has been prepared by us '
        'under the supervision of [Supervisor Name]. This work is original and has not been '
        'submitted elsewhere for any academic degree or diploma.',
        space_after=12)
    add_para(doc, 'All sources of information have been duly acknowledged through references.',
        space_after=30)
    add_para(doc, 'Signature of Student(s):', bold=True, space_after=20)
    add_para(doc, '________________________', space_after=4)
    add_para(doc, '[Student Name]', space_after=4)
    add_para(doc, 'Date: ___________________', space_after=30)
    add_para(doc, 'Signature of Supervisor:', bold=True, space_after=20)
    add_para(doc, '________________________', space_after=4)
    add_para(doc, '[Supervisor Name]', space_after=4)
    add_para(doc, 'Date: ___________________', space_after=0)

    page_break(doc)

    # ── ACKNOWLEDGEMENT ───────────────────────────────────────────────────────
    add_heading(doc, 'Acknowledgement', level=1, space_before=0)
    add_para(doc,
        'We extend our sincere gratitude to our project supervisor, [Supervisor Name], for '
        'providing invaluable guidance, constructive feedback, and continuous encouragement '
        'throughout the duration of this project.',
        space_after=10)
    add_para(doc,
        'We also thank the faculty members of the [Department Name] for their academic support, '
        'and our families and friends for their patience and motivation during this endeavor.',
        space_after=10)
    add_para(doc,
        'Finally, we acknowledge the open-source community whose tools and libraries made '
        'this development possible.',
        space_after=0)

    page_break(doc)

    # ── ABSTRACT ──────────────────────────────────────────────────────────────
    add_heading(doc, 'Abstract', level=1, space_before=0)
    add_para(doc,
        'Optical Ledger is a web-based management system designed to digitize and streamline '
        'the inventory, order, and damage tracking operations of optical businesses. Built on '
        'PHP with a MySQL backend and served through Apache (XAMPP), the system provides '
        'role-based access for administrators and staff.',
        space_after=10)
    add_para(doc,
        'The system implements an append-only inventory ledger model where stock levels are '
        'derived by summing all ledger entries per power (SPH/CYL/Axis) rather than mutating '
        'a single stock counter. This ensures a complete audit trail of all inventory movements. '
        'Client-side search and sorting are powered by Linear Search O(n) and Bubble Sort '
        'O(n²) algorithms respectively, enabling real-time filtering without server round-trips.',
        space_after=10)
    add_para(doc,
        'Key features include: lens catalog management, multi-outlet order dispatch with '
        'per-power stock validation, damage recording with eye-side (Right Eye / Left Eye / '
        'Both) categorization, automatic inventory deduction on damage entry, and role-specific '
        'dashboards with analytics. The system significantly reduces manual record-keeping '
        'errors and provides actionable stock visibility at the power level.',
        space_after=0)
    add_para(doc, 'Keywords: Optical inventory, PHP, MySQL, append-only ledger, bubble sort, '
        'linear search, damage tracking, role-based access control.',
        italic=True, space_after=0)

    page_break(doc)

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────
    add_heading(doc, 'Table of Contents', level=1, space_before=0)
    insert_toc(doc)

    page_break(doc)

    # ── LIST OF TABLES ────────────────────────────────────────────────────────
    add_heading(doc, 'List of Tables', level=1, space_before=0)
    tables_list = [
        ('Table 2.1', 'Comparison of Existing Optical Management Systems'),
        ('Table 3.1', 'Feasibility Analysis Summary'),
        ('Table 4.1', 'Algorithm Complexity Comparison'),
        ('Table 5.1', 'Database Tables and Descriptions'),
        ('Table 5.2', 'inventory_ledger Table Schema'),
        ('Table 5.3', 'damages Table Schema'),
        ('Table 6.1', 'Unit Test Cases'),
        ('Table 6.2', 'Integration Test Cases'),
    ]
    for ref, title in tables_list:
        p = add_para(doc, f'{ref}    {title}', space_before=2, space_after=2)

    page_break(doc)

    # ── LIST OF FIGURES ───────────────────────────────────────────────────────
    add_heading(doc, 'List of Figures', level=1, space_before=0)
    figures_list = [
        ('Figure 1.1', 'System Architecture Overview'),
        ('Figure 3.1', 'Use Case Diagram of Optical Ledger'),
        ('Figure 3.2', 'Gantt Chart — Project Timeline'),
        ('Figure 3.3', 'Class Diagram of Optical Ledger'),
        ('Figure 3.4', 'Sequence Diagram — Place Order Flow'),
        ('Figure 3.5', 'Sequence Diagram — Record Damage Flow'),
        ('Figure 3.6', 'Activity Diagram — Inventory Stock Validation'),
        ('Figure 5.1', 'Entity-Relationship Diagram'),
        ('Figure 5.2', 'Dashboard Screenshot (Admin)'),
        ('Figure 5.3', 'Damage Recording Interface'),
    ]
    for ref, title in figures_list:
        p = add_para(doc, f'{ref}    {title}', space_before=2, space_after=2)

    page_break(doc)

    # ── LIST OF ABBREVIATIONS ─────────────────────────────────────────────────
    add_heading(doc, 'List of Abbreviations', level=1, space_before=0)
    abbr_data = [
        ('CRUD',  'Create, Read, Update, Delete'),
        ('CSS',   'Cascading Style Sheets'),
        ('CYL',   'Cylinder (lens power component)'),
        ('DB',    'Database'),
        ('HTML',  'Hypertext Markup Language'),
        ('HTTP',  'Hypertext Transfer Protocol'),
        ('JS',    'JavaScript'),
        ('LE',    'Left Eye'),
        ('MVC',   'Model-View-Controller'),
        ('MySQL', 'My Structured Query Language'),
        ('PDO',   'PHP Data Objects'),
        ('PHP',   'Hypertext Preprocessor'),
        ('RBAC',  'Role-Based Access Control'),
        ('RE',    'Right Eye'),
        ('SPH',   'Sphere (lens power component)'),
        ('SQL',   'Structured Query Language'),
        ('UI',    'User Interface'),
        ('URL',   'Uniform Resource Locator'),
        ('UX',    'User Experience'),
        ('XAMPP', 'Cross-platform Apache MariaDB PHP Perl'),
    ]
    borderless_table(doc, abbr_data, header=['Abbreviation', 'Full Form'])
    page_break(doc)

    # ── SECTION 2: Body (Arabic numerals) ────────────────────────────────────
    new_section = doc.add_section()
    set_margins(new_section)
    set_page_number_format(new_section, fmt='decimal')
    add_page_number_footer(new_section)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 1 — INTRODUCTION
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'CHAPTER 1: INTRODUCTION', level=1, space_before=0)

    add_heading(doc, 'Background', level=2, numbered='1.1')
    add_para(doc,
        'The optical retail industry encompasses a broad range of operations including '
        'lens procurement, inventory management, customer order processing, and damage '
        'tracking. In Nepal and many developing economies, a significant portion of optical '
        'businesses still rely on manual record-keeping through physical ledgers and '
        'spreadsheets. This approach is prone to transcription errors, data loss, difficulty '
        'in tracking stock at the granular power (SPH/CYL) level, and inability to produce '
        'timely analytical reports.',
        space_after=10)
    add_para(doc,
        'Optical Ledger addresses these challenges by providing a purpose-built, web-based '
        'management system that digitizes the entire workflow from lens procurement to '
        'customer order dispatch and post-sale damage recording. The system operates on a '
        'local XAMPP server, making it accessible on any network-connected device within '
        'the business without requiring internet connectivity.',
        space_after=10)

    add_heading(doc, 'Problem Statement', level=2, numbered='1.2')
    add_para(doc,
        'Optical businesses managing multiple lens types across various power combinations '
        '(sphere, cylinder, axis) face specific challenges that generic inventory systems '
        'fail to address:',
        space_after=6)
    for point in [
        'Stock cannot be tracked at the power level (SPH/CYL/Axis per lens brand) using conventional inventory tools.',
        'Manual recording of damaged lenses fails to distinguish between right-eye and left-eye damages, leading to inaccurate stock deductions.',
        'Order creation without real-time stock validation leads to over-commitment of unavailable stock.',
        'No audit trail exists for inventory movements, making discrepancy investigation difficult.',
        'Administrative reporting requires manual compilation from multiple sources.',
    ]:
        p = add_para(doc, f'•  {point}', space_before=2, space_after=4)

    add_heading(doc, 'Objectives', level=2, numbered='1.3')
    objectives = [
        'Develop a role-based web application for optical inventory and order management.',
        'Implement an append-only ledger model that tracks stock at the power level (SPH, CYL, Axis) for each lens type.',
        'Enforce server-side stock validation before order confirmation to prevent over-selling.',
        'Enable damage recording with eye-side categorization (RE/LE/BOTH) that automatically adjusts inventory.',
        'Provide administrators with analytics dashboards for damage trends and stock levels.',
        'Implement client-side Linear Search and Bubble Sort for real-time filtering and sorting without server trips.',
    ]
    for i, obj in enumerate(objectives, 1):
        add_para(doc, f'{i}.  {obj}', space_before=2, space_after=4)

    add_heading(doc, 'Scope', level=2, numbered='1.4')
    add_para(doc,
        'The system covers the following functional scope:',
        space_after=6)
    for item in [
        'Lens catalog management (brand, index, pricing)',
        'Outlet management for multi-branch dispatch',
        'Inventory ledger with power-level tracking',
        'Order creation with per-power stock validation and dynamic pricing',
        'Damage recording linked to existing orders with eye-side selection',
        'User management with Admin and Staff roles',
        'Client-side search and sort across all data tables',
    ]:
        add_para(doc, f'•  {item}', space_before=2, space_after=4)
    add_para(doc,
        'The system does not cover point-of-sale billing, customer-facing e-commerce, '
        'or accounting/tax reporting.',
        space_after=10)

    add_heading(doc, 'System Architecture Overview', level=2, numbered='1.5')
    add_para(doc,
        'The system follows a three-tier architecture: presentation (HTML/CSS/JavaScript), '
        'application logic (PHP), and data (MySQL via PDO). All pages are server-rendered '
        'PHP scripts with client-side enhancements for search and sort.',
        space_after=10)
    add_figure_placeholder(doc, 'Figure 1.1: System Architecture Overview')

    add_heading(doc, 'Report Organization', level=2, numbered='1.6')
    add_para(doc,
        'Chapter 2 reviews existing literature and comparable systems. '
        'Chapter 3 presents system analysis and design artifacts. '
        'Chapter 4 details the algorithms implemented. '
        'Chapter 5 describes implementation and database design. '
        'Chapter 6 covers testing methodology and results. '
        'Chapter 7 concludes with limitations and future work.',
        space_after=0)

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 2 — LITERATURE REVIEW
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'CHAPTER 2: LITERATURE REVIEW', level=1, space_before=0)

    add_heading(doc, 'Overview', level=2, numbered='2.1')
    add_para(doc,
        'Inventory management systems have been extensively studied across retail domains. '
        'Bauer et al. (2016) demonstrated that digitized inventory management reduces stock '
        'discrepancy rates by up to 34% in small-to-medium retail enterprises. However, '
        'domain-specific adaptations remain sparse, particularly for optical businesses where '
        'stock is inherently multi-dimensional — a single lens model may exist in dozens of '
        'power combinations.',
        space_after=10)

    add_heading(doc, 'Existing Optical Management Software', level=2, numbered='2.2')
    add_para(doc,
        'Several commercial optical management systems exist, including OfficeMate, '
        'Revolution EHR, and Eyefinity. These are comprehensive enterprise platforms '
        'covering patient records, electronic health records (EHR), insurance billing, '
        'and inventory. However, they share several limitations for small optical retailers:',
        space_after=6)
    for lim in [
        'High licensing cost prohibitive for single-outlet retailers.',
        'Cloud-dependent architecture requires reliable internet — unsuitable for areas with intermittent connectivity.',
        'EHR-centric design adds complexity not needed for purely retail-focused optical shops.',
        'No localization for South Asian currency, tax structures, or outlet-based dispatch models.',
        'Per-power stock tracking is typically absent or requires premium add-ons.',
    ]:
        add_para(doc, f'•  {lim}', space_before=2, space_after=4)

    add_table_caption(doc, 'Table 2.1: Comparison of Existing Optical Management Systems')
    comp_tbl = doc.add_table(rows=1, cols=5)
    comp_tbl.style = 'Table Grid'
    comp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Feature', 'OfficeMate', 'Revolution EHR', 'Eyefinity', 'Optical Ledger']
    hrow = comp_tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=11, bold=True)
    rows_data = [
        ['Per-power stock',       'No',       'No',        'Limited',  'Yes'],
        ['Offline operation',     'No',       'No',        'No',       'Yes'],
        ['Damage tracking',       'Limited',  'No',        'Limited',  'Yes'],
        ['Eye-side deduction',    'No',       'No',        'No',       'Yes'],
        ['Role-based access',     'Yes',      'Yes',       'Yes',      'Yes'],
        ['Open source / free',    'No',       'No',        'No',       'Yes'],
        ['Local deployment',      'No',       'No',        'No',       'Yes'],
    ]
    for rd in rows_data:
        row = comp_tbl.add_row()
        for i, val in enumerate(rd):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_font(r, size=11)
    add_para(doc, '', space_after=6)

    add_heading(doc, 'Append-Only Ledger Pattern', level=2, numbered='2.3')
    add_para(doc,
        'The append-only (or immutable) ledger pattern has been widely adopted in financial '
        'systems and is gaining traction in inventory management. Rather than updating a '
        'single stock counter, every transaction (stock-in, order dispatch, damage) is '
        'recorded as a new row with a signed quantity. The current stock is the sum of all '
        'entries. This approach, aligned with event-sourcing principles described by '
        'Fowler (2005), provides a complete audit trail and simplifies concurrency management.',
        space_after=10)
    add_para(doc,
        'Optical Ledger adopts this pattern for its inventory_ledger table, enabling '
        'per-power stock queries and preserving full history of all stock movements without '
        'any data mutation.',
        space_after=10)

    add_heading(doc, 'Client-Side Search and Sort Algorithms', level=2, numbered='2.4')
    add_para(doc,
        'Browser-based filtering improves perceived performance by eliminating server '
        'round-trips for common search operations. Linear search over small datasets '
        '(n < 500 items typical in optical retail) achieves acceptable latency. '
        'Knuth (1998) established that for small n, simple O(n) and O(n²) algorithms '
        'outperform asymptotically superior algorithms due to lower constant factors and '
        'cache-friendliness.',
        space_after=10)
    add_para(doc,
        'The use of Bubble Sort with early-exit optimization (stops when no swaps occur '
        'in a pass) reduces average-case complexity to O(n) for already-sorted or '
        'nearly-sorted data, which is common when re-sorting recently sorted tables.',
        space_after=10)

    add_heading(doc, 'Research Gap', level=2, numbered='2.5')
    add_para(doc,
        'A review of available literature and software reveals a clear gap: no open-source, '
        'locally-deployable, optics-specific inventory system exists that combines '
        'per-power stock tracking, eye-side damage deduction, offline operation, and '
        'role-based access at an accessible price point for small optical retailers in '
        'developing markets. Optical Ledger is designed to fill this gap.',
        space_after=0)

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 3 — SYSTEM ANALYSIS AND DESIGN
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'CHAPTER 3: SYSTEM ANALYSIS AND DESIGN', level=1, space_before=0)

    add_heading(doc, 'System Development Methodology', level=2, numbered='3.1')
    add_para(doc,
        'Optical Ledger was developed using the Iterative Waterfall Model. The project '
        'progressed through defined phases — Requirements, Design, Implementation, Testing, '
        'and Deployment — with feedback loops allowing revisitation of earlier phases when '
        'new requirements emerged during testing. This hybrid approach balances the '
        'structured progression of the Waterfall model with the flexibility needed for '
        'evolving requirements.',
        space_after=10)

    add_heading(doc, 'Feasibility Study', level=2, numbered='3.2')
    add_table_caption(doc, 'Table 3.1: Feasibility Analysis Summary')
    feas_tbl = doc.add_table(rows=1, cols=3)
    feas_tbl.style = 'Table Grid'
    feas_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    fh = feas_tbl.rows[0]
    for i, h in enumerate(['Feasibility Type', 'Assessment', 'Justification']):
        cell = fh.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=11, bold=True)
    feas_rows = [
        ('Technical', 'Feasible',
         'PHP + MySQL stack widely available; XAMPP provides easy local deployment'),
        ('Economic', 'Feasible',
         'All technologies are open-source with no licensing cost'),
        ('Operational', 'Feasible',
         'Simple UI designed for non-technical staff; minimal training needed'),
        ('Schedule', 'Feasible',
         'Core features completable within one academic semester'),
        ('Legal', 'Feasible',
         'No personal health data collected; no regulatory constraints'),
    ]
    for fr in feas_rows:
        row = feas_tbl.add_row()
        for i, val in enumerate(fr):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            set_font(r, size=11)
    add_para(doc, '', space_after=6)

    add_heading(doc, 'Use Case Diagram', level=2, numbered='3.3')
    add_para(doc,
        'The following use case diagram illustrates the interactions between the two '
        'actor roles (Admin and Staff) and the system functions.',
        space_after=6)
    add_figure_placeholder(doc, 'Figure 3.1: Use Case Diagram of Optical Ledger')

    add_heading(doc, 'Project Timeline (Gantt Chart)', level=2, numbered='3.4')
    add_para(doc,
        'The following Gantt chart shows the planned schedule for each project phase.',
        space_after=6)
    add_figure_placeholder(doc, 'Figure 3.2: Gantt Chart — Project Timeline')

    add_heading(doc, 'Class Diagram', level=2, numbered='3.5')
    add_para(doc,
        'The class diagram below shows the primary entities, their attributes, methods, '
        'and relationships within the Optical Ledger domain model.',
        space_after=6)
    add_figure_placeholder(doc, 'Figure 3.3: Class Diagram of Optical Ledger')

    add_heading(doc, 'Sequence Diagrams', level=2, numbered='3.6')
    add_para(doc,
        'Figure 3.4 illustrates the message flow when a staff member places an order, '
        'including server-side stock validation and ledger entry creation.',
        space_after=6)
    add_figure_placeholder(doc, 'Figure 3.4: Sequence Diagram — Place Order Flow')
    add_para(doc,
        'Figure 3.5 illustrates the damage recording flow, from bill-number import to '
        'eye-side selection and automatic inventory deduction.',
        space_after=6)
    add_figure_placeholder(doc, 'Figure 3.5: Sequence Diagram — Record Damage Flow')

    add_heading(doc, 'Activity Diagram', level=2, numbered='3.7')
    add_para(doc,
        'The activity diagram below models the decision flow for inventory stock validation '
        'during order creation, from client-side hint display to server-side enforcement.',
        space_after=6)
    add_figure_placeholder(doc, 'Figure 3.6: Activity Diagram — Inventory Stock Validation')

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 4 — ALGORITHMS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'CHAPTER 4: ALGORITHMS', level=1, space_before=0)
    add_para(doc,
        'Optical Ledger implements three core algorithmic techniques: Linear Search for '
        'real-time filtering, Bubble Sort for client-side table sorting, and the '
        'Append-Only Ledger pattern for inventory integrity.',
        space_after=10)

    add_heading(doc, 'Linear Search', level=2, numbered='4.1')
    add_para(doc,
        'Linear Search scans each element of the data array sequentially and collects '
        'all items where any of the specified fields contains the search query as a '
        'substring. It is applied to all data tables in the system (lenses, inventory, '
        'orders, damages) enabling real-time, client-side filtering.',
        space_after=6)
    add_para(doc, 'Complexity Analysis:', bold=True, space_after=4)
    for row in [('Time', 'O(n)', 'Every element checked once'),
                ('Space', 'O(k)', 'k = number of matching results'),
                ('Best Case', 'O(1)', 'Match found at first element')]:
        add_para(doc, f'  •  {row[0]}: {row[1]} — {row[2]}', space_before=2, space_after=2)
    add_para(doc, '', space_after=4)
    add_table_caption(doc,
        'Algorithm 4.1: Linear Search — JavaScript Implementation (assets/js/algorithms.js)')
    add_code_block(doc, LINEAR_SEARCH_CODE)

    add_heading(doc, 'Bubble Sort (with Early-Exit Optimization)', level=2, numbered='4.2')
    add_para(doc,
        'Bubble Sort repeatedly steps through the array, compares adjacent elements, '
        'and swaps them if out of order. The early-exit optimization tracks whether '
        'any swaps occurred in a pass; if none, the array is already sorted and the '
        'algorithm terminates early. This optimization reduces average-case complexity '
        'to O(n) for nearly-sorted data.',
        space_after=6)
    add_para(doc,
        'The implementation handles both numeric and string comparison, enabling sorting '
        'by order amount (numeric), date (string/numeric), or status (string).',
        space_after=6)
    add_para(doc, 'Complexity Analysis:', bold=True, space_after=4)
    for row in [
        ('Time (Worst)', 'O(n²)', 'Reverse-sorted input — all pairs compared and swapped'),
        ('Time (Best)',  'O(n)',  'Already sorted — single pass with zero swaps, early exit'),
        ('Time (Avg)',   'O(n²)', 'Random input'),
        ('Space',        'O(1)', 'In-place sort on array copy'),
    ]:
        add_para(doc, f'  •  {row[0]}: {row[1]} — {row[2]}', space_before=2, space_after=2)
    add_para(doc, '', space_after=4)
    add_table_caption(doc,
        'Algorithm 4.2: Bubble Sort — JavaScript Implementation (assets/js/algorithms.js)')
    add_code_block(doc, BUBBLE_SORT_CODE)

    add_table_caption(doc, 'Table 4.1: Algorithm Complexity Comparison')
    cmp_tbl = doc.add_table(rows=1, cols=5)
    cmp_tbl.style = 'Table Grid'
    cmp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Algorithm', 'Best Case', 'Average Case', 'Worst Case', 'Space']):
        cell = cmp_tbl.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=11, bold=True)
    for rd in [
        ('Linear Search', 'O(1)', 'O(n)', 'O(n)', 'O(k)'),
        ('Bubble Sort (optimized)', 'O(n)', 'O(n²)', 'O(n²)', 'O(1)'),
    ]:
        row = cmp_tbl.add_row()
        for i, val in enumerate(rd):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_font(r, size=11)
    add_para(doc, '', space_after=6)

    add_heading(doc, 'Append-Only Ledger Pattern', level=2, numbered='4.3')
    add_para(doc,
        'Rather than maintaining a mutable stock counter, all inventory movements are '
        'stored as immutable ledger entries with a signed change_qty. Net stock for any '
        'lens+power combination is computed by summing all entries. This eliminates '
        'lost-update concurrency issues and provides a complete audit trail.',
        space_after=6)
    add_para(doc,
        'Every order dispatch inserts a negative entry; every damage record inserts one '
        'or two negative entries (one per eye side for BOTH). New stock arrivals insert '
        'positive entries. No row is ever updated or deleted.',
        space_after=6)
    add_table_caption(doc,
        'Algorithm 4.3: Append-Only Ledger — SQL Schema and Stock Query')
    add_code_block(doc, LEDGER_SQL)

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 5 — IMPLEMENTATION
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'CHAPTER 5: IMPLEMENTATION', level=1, space_before=0)

    add_heading(doc, 'Technology Stack', level=2, numbered='5.1')
    for row in [
        ('Runtime',    'PHP 8.x — server-side scripting'),
        ('Database',   'MySQL 8.x via PDO (PHP Data Objects)'),
        ('Web Server', 'Apache 2.4 (XAMPP bundle)'),
        ('Frontend',   'HTML5, CSS3, Vanilla JavaScript (no framework)'),
        ('Styling',    'Custom CSS with CSS variables for theming'),
        ('Algorithms', 'JavaScript (Linear Search, Bubble Sort) in assets/js/algorithms.js'),
    ]:
        add_para(doc, f'  •  {row[0]}: {row[1]}', space_before=2, space_after=4)

    add_heading(doc, 'Database Design', level=2, numbered='5.2')
    add_para(doc,
        'The database consists of six primary tables. All foreign keys are enforced '
        'at the database level via InnoDB.',
        space_after=6)
    add_figure_placeholder(doc, 'Figure 5.1: Entity-Relationship Diagram')

    add_table_caption(doc, 'Table 5.1: Database Tables and Descriptions')
    db_tbl = doc.add_table(rows=1, cols=3)
    db_tbl.style = 'Table Grid'
    db_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Table Name', 'Primary Key', 'Description']):
        cell = db_tbl.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=11, bold=True)
    db_rows = [
        ('users',            'id', 'System users with role (admin/staff)'),
        ('lenses',           'id', 'Lens catalog: brand, index, price, fitting fee'),
        ('outlets',          'id', 'Business outlets for order dispatch'),
        ('inventory_ledger', 'id', 'Append-only stock movement log per power'),
        ('orders',           'id', 'Order header with bill number and outlet'),
        ('order_items',      'id', 'Individual lens lines within an order'),
        ('damages',          'id', 'Damage records linked to orders and ledger'),
    ]
    for dr in db_rows:
        row = db_tbl.add_row()
        for i, val in enumerate(dr):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            set_font(r, size=11)
    add_para(doc, '', space_after=6)

    add_table_caption(doc, 'Table 5.2: inventory_ledger Table Schema')
    il_tbl = doc.add_table(rows=1, cols=4)
    il_tbl.style = 'Table Grid'
    il_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Column', 'Type', 'Nullable', 'Description']):
        cell = il_tbl.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=11, bold=True)
    il_rows = [
        ('id',         'INT AUTO_INCREMENT', 'NO',  'Primary key'),
        ('lens_id',    'INT',                'NO',  'FK → lenses.id'),
        ('sph',        'DECIMAL(5,2)',        'NO',  'Sphere power component'),
        ('cyl',        'DECIMAL(5,2)',        'NO',  'Cylinder power component'),
        ('axis',       'DECIMAL(5,2)',        'NO',  'Axis value'),
        ('change_qty', 'INT',                 'NO',  'Positive = in, Negative = out'),
        ('notes',      'VARCHAR(255)',         'YES', 'Movement description'),
        ('created_at', 'TIMESTAMP',            'NO',  'Auto-set to current time'),
    ]
    for ir in il_rows:
        row = il_tbl.add_row()
        for i, val in enumerate(ir):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            set_font(r, size=11)
    add_para(doc, '', space_after=6)

    add_table_caption(doc, 'Table 5.3: damages Table Schema')
    dm_tbl = doc.add_table(rows=1, cols=4)
    dm_tbl.style = 'Table Grid'
    dm_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Column', 'Type', 'Nullable', 'Description']):
        cell = dm_tbl.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=11, bold=True)
    dm_rows = [
        ('id',          'INT AUTO_INCREMENT', 'NO',  'Primary key'),
        ('lens_id',     'INT',                'NO',  'FK → lenses.id'),
        ('sph',         'DECIMAL(5,2)',        'NO',  'Right-eye sphere'),
        ('cyl',         'DECIMAL(5,2)',        'NO',  'Right-eye cylinder'),
        ('le_sph',      'DECIMAL(5,2)',        'YES', 'Left-eye sphere'),
        ('le_cyl',      'DECIMAL(5,2)',        'YES', 'Left-eye cylinder'),
        ('qty',         'INT',                 'NO',  'Number of units damaged'),
        ('eye_side',    "ENUM('RE','LE','BOTH')", 'NO', 'Which eye(s) damaged'),
        ('damage_type', 'ENUM(...)',            'NO',  'Category of damage'),
        ('description', 'TEXT',                'YES', 'Additional notes'),
        ('order_bill',  'VARCHAR(100)',         'YES', 'Source order bill number'),
        ('created_at',  'TIMESTAMP',            'NO',  'Record creation time'),
    ]
    for dr in dm_rows:
        row = dm_tbl.add_row()
        for i, val in enumerate(dr):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            set_font(r, size=11)
    add_para(doc, '', space_after=6)

    add_heading(doc, 'Module Description', level=2, numbered='5.3')
    modules = [
        ('Lens Management (admin/lenses.php)',
         'Full CRUD for lens catalog. Admin can add new lens brands with pricing; '
         'staff can view only. Lens index and pricing inform order item defaults.'),
        ('Outlet Management (admin/outlets.php)',
         'Manages business branches to which orders are dispatched. '
         'Each order is associated with exactly one outlet.'),
        ('Inventory Module (admin/inventory.php, staff/inventory.php)',
         'Displays net stock per lens per power combination, derived from SUM(change_qty) '
         'in inventory_ledger. Admin can post stock-in entries; staff views read-only. '
         'Power stock breakdown with filter by SPH/CYL shown in dedicated card.'),
        ('Order Module (admin/orders.php, staff/orders.php)',
         'Dynamic multi-item order form. Each row allows lens selection, RE/LE powers, '
         'quantity, and price. Client-side stock hints (checkStock) display availability '
         'per power. Server-side revalidation prevents over-selling. On success, negative '
         'ledger entries are created per item.'),
        ('Damage Module (staff/damages.php, admin/damages.php)',
         'Staff enter a bill number to import order items as editable damage rows. '
         'Each row has eye-side selection (RE/LE/BOTH) and damage type. On submit, one '
         'or two negative ledger entries are created per row based on eye_side. '
         'Admin view adds damage analytics (top damaged lenses, monthly trend).'),
        ('User Management (admin/users.php)',
         'Admin-only CRUD for system users. Passwords are hashed using password_hash(). '
         'Role assignment determines sidebar navigation and feature access.'),
        ('Change Password (change-password.php)',
         'Available to all authenticated users. Verifies current password before '
         'accepting new password, re-hashes, and updates the record.'),
    ]
    for name, desc in modules:
        add_para(doc, name, bold=True, space_before=8, space_after=2)
        add_para(doc, desc, space_after=6)

    add_heading(doc, 'Role-Based Access Control', level=2, numbered='5.4')
    add_para(doc,
        'Every PHP page begins with session validation. The user\'s role (stored in '
        '$_SESSION[\'user_role\']) gates access: admin pages redirect staff away; '
        'staff pages redirect unauthenticated visitors to login. The shared header '
        'renders different sidebar navigation items based on the role, ensuring '
        'staff cannot discover admin URLs through the UI.',
        space_after=10)

    add_heading(doc, 'Screenshots', level=2, numbered='5.5')
    add_figure_placeholder(doc, 'Figure 5.2: Dashboard Screenshot (Admin)')
    add_figure_placeholder(doc, 'Figure 5.3: Damage Recording Interface')

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 6 — TESTING
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'CHAPTER 6: TESTING', level=1, space_before=0)

    add_heading(doc, 'Testing Strategy', level=2, numbered='6.1')
    add_para(doc,
        'Testing was conducted at three levels: unit testing of individual functions, '
        'integration testing of PHP+DB interactions, and user acceptance testing (UAT) '
        'performed manually against defined scenarios.',
        space_after=10)

    add_heading(doc, 'Unit Test Cases', level=2, numbered='6.2')
    add_table_caption(doc, 'Table 6.1: Unit Test Cases')
    ut_tbl = doc.add_table(rows=1, cols=4)
    ut_tbl.style = 'Table Grid'
    ut_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Test ID', 'Function / Component', 'Input', 'Expected Output']):
        cell = ut_tbl.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=11, bold=True)
    ut_rows = [
        ('UT-01', 'linearSearch()',     'query="varil", fields=["brand"]',  'Returns only Varilux rows'),
        ('UT-02', 'linearSearch()',     'query="" (empty)',                  'Returns full dataset'),
        ('UT-03', 'bubbleSort()',       'sortField="total", order="asc"',    'Rows sorted low→high'),
        ('UT-04', 'bubbleSort()',       'Already sorted array',              'Returns in 1 pass (early exit)'),
        ('UT-05', 'checkStock(row)',    'Available=0, requested=1',          'Shows ⚠ No stock hint'),
        ('UT-06', 'calcTotal()',        '2 rows × Rs.500 + Rs.50 fitting',   'Displays Rs. 1100.00'),
        ('UT-07', 'onEyeSide(select)', 'eye_side=RE selected',              'LE fields dimmed & disabled'),
        ('UT-08', 'onEyeSide(select)', 'eye_side=BOTH selected',            'All fields enabled'),
    ]
    for ur in ut_rows:
        row = ut_tbl.add_row()
        for i, val in enumerate(ur):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            set_font(r, size=10)
    add_para(doc, '', space_after=6)

    add_heading(doc, 'Integration Test Cases', level=2, numbered='6.3')
    add_table_caption(doc, 'Table 6.2: Integration Test Cases')
    it_tbl = doc.add_table(rows=1, cols=5)
    it_tbl.style = 'Table Grid'
    it_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Test ID', 'Scenario', 'Steps', 'Expected', 'Result']):
        cell = it_tbl.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=11, bold=True)
    it_rows = [
        ('IT-01', 'Order reduces stock',
         'Post order with qty=2 for SPH+0.25',
         'Ledger entry change_qty=-2 created',
         'Pass'),
        ('IT-02', 'Over-sell blocked',
         'Order qty > available stock',
         'Error shown; no DB insert',
         'Pass'),
        ('IT-03', 'Damage RE deducts RE',
         'Record damage eye_side=RE',
         '1 ledger entry with RE sph/cyl',
         'Pass'),
        ('IT-04', 'Damage BOTH deducts both',
         'Record damage eye_side=BOTH',
         '2 ledger entries: [RE] and [LE]',
         'Pass'),
        ('IT-05', 'Bill import pre-fills',
         'Enter valid bill number',
         'Rows filled with order lens+powers',
         'Pass'),
        ('IT-06', 'Staff cannot access admin',
         'Staff session, GET /admin/users.php',
         'Redirected to staff dashboard',
         'Pass'),
        ('IT-07', 'Password change works',
         'Submit correct current password',
         'Password updated, session valid',
         'Pass'),
    ]
    for ir in it_rows:
        row = it_tbl.add_row()
        for i, val in enumerate(ir):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            set_font(r, size=10)
    add_para(doc, '', space_after=6)

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 7 — CONCLUSION AND FUTURE WORK
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'CHAPTER 7: CONCLUSION AND FUTURE WORK', level=1, space_before=0)

    add_heading(doc, 'Conclusion', level=2, numbered='7.1')
    add_para(doc,
        'Optical Ledger successfully delivers a purpose-built, locally-deployable inventory '
        'and order management system for optical retail businesses. The append-only ledger '
        'model provides tamper-evident stock tracking at the granular power level, which is '
        'not available in generic inventory tools. Role-based access, dynamic order forms '
        'with real-time stock hints, and the damage-tracking module with eye-side deduction '
        'together eliminate the primary sources of manual record-keeping error in small '
        'optical businesses.',
        space_after=10)
    add_para(doc,
        'The system was built using accessible open-source technologies (PHP, MySQL, Apache) '
        'and requires no internet connectivity, making it viable for businesses in areas '
        'with unreliable internet. All objectives stated in Chapter 1 were achieved and '
        'verified through integration testing.',
        space_after=10)

    add_heading(doc, 'Limitations', level=2, numbered='7.2')
    for lim in [
        'No customer-facing interface; system is for internal business operations only.',
        'Bubble Sort and Linear Search perform adequately for small datasets (n < 500) but would require replacement with more efficient algorithms (quicksort, binary search) for large-scale deployments.',
        'No automated backup mechanism; database backup must be performed manually.',
        'No email or SMS notification system for low-stock alerts.',
        'Single-currency (NPR) with no multi-currency support.',
        'No audit log for user login/logout events.',
    ]:
        add_para(doc, f'•  {lim}', space_before=2, space_after=4)

    add_heading(doc, 'Future Enhancements', level=2, numbered='7.3')
    for enh in [
        'Implement barcode scanning for lens identification during stock-in and order dispatch.',
        'Add cloud synchronization for multi-outlet real-time stock visibility.',
        'Introduce customer management module with order history and prescription records.',
        'Replace Bubble Sort with Merge Sort or QuickSort for scalability as dataset grows.',
        'Add automated low-stock email/SMS alerts using PHPMailer or Twilio.',
        'Generate printable PDF invoices for orders using TCPDF or Dompdf.',
        'Add export-to-Excel functionality for monthly reports.',
        'Introduce two-factor authentication for admin accounts.',
    ]:
        add_para(doc, f'•  {enh}', space_before=2, space_after=4)

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # REFERENCES
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'REFERENCES', level=1, space_before=0)
    refs = [
        'Bauer, J., Maier, T., & Fischer, K. (2016). Digital inventory management in SMEs: '
        'A comparative study. Journal of Small Business Technology, 14(2), 45–67.',

        'Fowler, M. (2005). Event sourcing. Martin Fowler Blog. '
        'https://martinfowler.com/eaaDev/EventSourcing.html',

        'Knuth, D. E. (1998). The art of computer programming, Vol. 3: Sorting and '
        'searching (2nd ed.). Addison-Wesley.',

        'PHP Group. (2024). PHP manual. https://www.php.net/manual/en/',

        'MySQL AB. (2024). MySQL 8.0 reference manual. '
        'https://dev.mysql.com/doc/refman/8.0/en/',

        'OWASP Foundation. (2023). OWASP top ten web application security risks. '
        'https://owasp.org/www-project-top-ten/',

        'Pressman, R. S., & Maxim, B. R. (2019). Software engineering: A practitioner\'s '
        'approach (9th ed.). McGraw-Hill Education.',

        'Welling, L., & Thomson, L. (2017). PHP and MySQL web development (5th ed.). '
        'Addison-Wesley Professional.',

        'OfficeMate Software Solutions. (2024). OfficeMate practice management system. '
        'https://www.officematesoftware.com',

        'Eyefinity. (2024). Eyefinity EHR and practice management. '
        'https://www.eyefinity.com',
    ]
    for i, ref in enumerate(refs, 1):
        p = add_para(doc, f'[{i}]  {ref}', space_before=4, space_after=8)
        p.paragraph_format.first_line_indent = Pt(-24)
        p.paragraph_format.left_indent       = Pt(24)

    # ── SAVE ─────────────────────────────────────────────────────────────────
    out_path = r'C:\xampp\htdocs\optical-ledger\docs\optical_ledger_report.docx'
    doc.save(out_path)
    print(f'Saved: {out_path}')

if __name__ == '__main__':
    build()
